import asyncio
import datetime
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.domain.extraction import Link, Page, ParsedDocument

from .base_extraction_service import BaseExtractionService


class ParseWordService(BaseExtractionService):
    def __init__(self, word_path: str):
        self.word_path = Path(word_path)

    async def parse_to_json(self) -> ParsedDocument:
        pages = await asyncio.to_thread(self._extract_content)
        return self._build_response(pages)

    def _extract_content(self) -> list[Page]:
        doc = Document(str(self.word_path))
        text_parts = []
        links = []

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text_parts.append(para.text)
                links.extend(self._extract_links(para))
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                text_parts.append(self._extract_table(table))
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            links.extend(self._extract_links(para))

        full_text = "\n".join(part for part in text_parts if part.strip())
        return [Page(page_number=1, text=full_text, links=links)]

    def _extract_links(self, paragraph: Paragraph) -> list[Link]:
        links = []
        for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
            r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if r_id and r_id in paragraph.part.rels:
                url = paragraph.part.rels[r_id].target_ref
                text = "".join(node.text or "" for node in hyperlink.xpath(".//w:t"))
                if url and text:
                    links.append(Link(text=text.strip(), url=url))
        return links

    def _extract_table(self, table: Table) -> str:
        lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = "\t".join(cells)
            if line.strip():
                lines.append(line)
        return "\n".join(lines)

    def _build_response(self, pages: list[Page]) -> ParsedDocument:
        return ParsedDocument(
            file_type=self.word_path.suffix.lower(),
            processed_at=datetime.datetime.now().isoformat() + "Z",
            pages=pages,
            processing_method="word_extract",
        )
